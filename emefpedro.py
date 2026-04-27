import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# --- CONFIGURAÇÕES DE ESTILO (PALETA AZUL COM CAMPOS DESTACADOS) ---
st.set_page_config(page_title="Gestão EMEF Pedro Caminoto", layout="wide")

st.markdown("""
    <style>
    /* Fundo geral e abas */
    .stApp { background-color: #f4f7f9; }
    .stTabs [data-baseweb="tab-list"] { background-color: #004a99; padding: 15px 20px 0px 20px; border-radius: 12px; gap: 25px; box-shadow: 0 4px 10px rgba(0,0,0,0.15); }
    .stTabs [data-baseweb="tab"] { height: 60px; color: rgba(255, 255, 255, 0.7) !important; background-color: transparent !important; font-size: 16px; font-weight: 500; transition: all 0.3s ease; padding: 0 15px; }
    .stTabs [data-baseweb="tab"]:hover { color: #ffffff !important; background-color: rgba(255, 255, 255, 0.1) !important; transform: translateY(-2px); }
    .stTabs [aria-selected="true"] { color: #ffffff !important; background-color: transparent !important; border-bottom: 4px solid #ffffff !important; font-weight: 700 !important; }
    
    /* Card de Conteúdo */
    div[data-testid="stVerticalBlock"] > div.stVerticalBlock { background-color: white; padding: 30px; border-radius: 0 0 15px 15px; box-shadow: 0 8px 20px rgba(0,0,0,0.05); border: 1px solid #e0e6ed; margin-top: -10px; }

    /* --- ESTILIZAÇÃO DOS CAMPOS DE ENTRADA (MUDANÇA SOLICITADA) --- */
    /* Campos de Texto, Data e Seleção */
    div[data-baseweb="input"], div[data-baseweb="select"], div[data-baseweb="popover"], .stMultiSelect {
        background-color: #eef4ff !important; /* Azul muito claro para contraste */
        border-radius: 8px !important;
        border: 1px solid #c2d6ff !important; /* Borda suave azul */
    }
    
    /* Estilo interno do input */
    input, select, textarea {
        background-color: #eef4ff !important;
        color: #002d5c !important; /* Texto azul escuro para leitura */
    }

    /* Borda de Destaque ao clicar/focar no campo */
    div[data-baseweb="input"]:focus-within, div[data-baseweb="select"]:focus-within {
        border: 2px solid #004a99 !important;
        box-shadow: 0 0 5px rgba(0, 74, 153, 0.2) !important;
    }

    /* Botões */
    div.stButton > button { background-color: #0056b3; color: white; border-radius: 8px; border: none; padding: 8px 25px; font-weight: 600; transition: 0.2s; }
    div.stButton > button:hover { background-color: #003d80; color: white; }
    .btn-excluir button { background-color: #bd2130 !important; }

    /* Expanders */
    .stExpander { border: 1px solid #dee2e6 !important; background-color: #fafbfc !important; border-radius: 10px !important; margin-bottom: 10px; }
    </style>
    """, unsafe_allow_html=True)

# --- BANCO DE DADOS ---
DB_FILE = "dados_professores.json"
HIST_FILE = "historico_lotes.json"

def load_data(file):
    if os.path.exists(file):
        with open(file, "r", encoding="utf-8") as f: return json.load(f)
    return {}

def save_data(file, data):
    with open(file, "w", encoding="utf-8") as f: json.dump(data, f, indent=4, ensure_ascii=False)

if 'profs' not in st.session_state: st.session_state.profs = load_data(DB_FILE)
if 'lotes' not in st.session_state: st.session_state.lotes = load_data(HIST_FILE)

# --- FUNÇÃO GERADORA DE DOCX ---
def add_styled_text(cell, text, bold=False, size=8):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)

def gerar_docx_modelo(lote_data):
    doc = Document()
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Inches(0.4), Inches(0.4)
    section.left_margin, section.right_margin = Inches(0.5), Inches(0.5)

    for p_id, p in lote_data['professores'].items():
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.add_run("MUNICÍPIO DE TEODORO SAMPAIO\n").bold = True
        h.add_run("CNPJ Nº. 44.951.515/0001-42\n").font.size = Pt(9)
        h.add_run("EMEF “PEDRO CAMINOTO”\n").bold = True
        h.add_run("R. PROFESSORA APARECIDA MARIA DE SOUZA, Nº. 1.700 – VILA FURLAN – FONE: (18) 3282-3533.\nemef.pcaminoto@pmteodorosampaio.sp.gov.br\nTEODORO SAMPAIO – SP").font.size = Pt(8)

        t_tit = doc.add_table(rows=2, cols=1); t_tit.style = 'Table Grid'
        add_styled_text(t_tit.cell(0,0), "CONTROLE DE SUPORTE PEDAGÓGICO", bold=True, size=10)
        add_styled_text(t_tit.cell(1,0), f"{lote_data['inicio_f']} A {lote_data['fim_f']}", bold=True, size=10)

        t_info = doc.add_table(rows=3, cols=2); t_info.style = 'Table Grid'
        t_info.cell(0,0).text = f"Professor: {p['nome'].upper()}"
        t_info.cell(1,0).text = f"RG: {p['rg']}"
        t_info.cell(1,1).text = f"Situação: {p['situacao']}"
        
        suporte_txt = ""
        for dia, ag in p['agenda'].items():
            dia_abrev = dia[:3].upper()
            if ag['HTPC'][0]: suporte_txt += f"HTPC: {dia_abrev}: {ag['HTPC'][0]}-{ag['HTPC'][1]}  "
            if ag['HSP1'][0]: suporte_txt += f"HSP: {dia_abrev}: {ag['HSP1'][0]}-{ag['HSP1'][1]}  "
            if ag['HSP2'][0]: suporte_txt += f"HSP: {dia_abrev}: {ag['HSP2'][0]}-{ag['HSP2'][1]}  "
            if ag['HE'][0]: suporte_txt += f"HE: {dia_abrev}: {ag['HE'][0]}-{ag['HE'][1]}  "
        
        t_info.cell(2,0).text = f"Hora de Suporte: {suporte_txt}"
        t_info.cell(2,1).text = f"Disciplina: {p['disciplina']}\nCategoria: {p['categoria']}"

        doc.add_paragraph()
        t_cal = doc.add_table(rows=2, cols=10); t_cal.style = 'Table Grid'
        add_styled_text(t_cal.cell(0,0).merge(t_cal.cell(1,0)), "DATA", True)
        add_styled_text(t_cal.cell(0,1).merge(t_cal.cell(0,2)), "HTPC", True)
        add_styled_text(t_cal.cell(0,3).merge(t_cal.cell(0,4)), "HSP", True)
        add_styled_text(t_cal.cell(0,5).merge(t_cal.cell(0,6)), "HSP", True)
        add_styled_text(t_cal.cell(0,7).merge(t_cal.cell(0,8)), "HE", True)
        add_styled_text(t_cal.cell(0,9).merge(t_cal.cell(1,9)), "VISTO", True)
        for i in range(1, 9): add_styled_text(t_cal.cell(1, i), "Início" if i % 2 != 0 else "Término", size=7)

        d_ini = datetime.strptime(lote_data['inicio'], "%Y-%m-%d")
        d_fim = datetime.strptime(lote_data['fim'], "%Y-%m-%d")
        curr = d_ini
        trad = {0:"SEGUNDA", 1:"TERÇA", 2:"QUARTA", 3:"QUINTA", 4:"SEXTA", 5:"SÁBADO", 6:"DOMINGO"}
        while curr <= d_fim:
            row = t_cal.add_row().cells
            add_styled_text(row[0], curr.strftime("%d"), size=8)
            wd = trad[curr.weekday()]
            if wd in ["SÁBADO", "DOMINGO"]:
                m = row[1].merge(row[8]); add_styled_text(m, wd, bold=True, size=8)
            else:
                cfg = p['agenda'].get(wd.capitalize(), {"HTPC":["",""],"HSP1":["",""],"HSP2":["",""],"HE":["",""]})
                add_styled_text(row[1], cfg["HTPC"][0]); add_styled_text(row[2], cfg["HTPC"][1])
                add_styled_text(row[3], cfg["HSP1"][0]); add_styled_text(row[4], cfg["HSP1"][1])
                add_styled_text(row[5], cfg["HSP2"][0]); add_styled_text(row[6], cfg["HSP2"][1])
                add_styled_text(row[7], cfg["HE"][0]); add_styled_text(row[8], cfg["HE"][1])
            curr += timedelta(days=1)

        doc.add_paragraph("\n\n")
        sig = doc.add_table(rows=1, cols=2)
        add_styled_text(sig.cell(0,0), "__________________________________\nCoordenador", size=9)
        add_styled_text(sig.cell(0,1), "__________________________________\nDiretor", size=9)
        if p_id != list(lote_data['professores'].keys())[-1]: doc.add_page_break()

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# --- INTERFACE ---
st.title("🏛️ Portal Pedro Caminoto - Gestão Pedagógica")
tabs = st.tabs(["📋 Cadastrar", "📝 Editar Cadastro", "🚀 Gerar Relatório", "📚 Histórico"])

# 1. ABA CADASTRAR
with tabs[0]:
    st.subheader("Novo Cadastro")
    with st.form("cad_form", clear_on_submit=True):
        c1, c2, c3 = st.columns(3); nome = c1.text_input("Nome Completo"); rg = c2.text_input("RG"); situacao = c3.selectbox("Situação", ["ACTS", "EFETIVO"])
        c4, c5, c6 = st.columns([2,2,1]); disciplina = c4.text_input("Disciplina Principal"); categoria = c5.text_input("Categoria"); ativo = c6.checkbox("Professor Ativo", value=True)
        st.write("---"); st.write("📌 **Horários Fixos**")
        agenda_nova = {}
        for d in ["Segunda", "Terça", "Quarta", "Quinta", "Sexta"]:
            with st.expander(f"📅 {d}"):
                cols = st.columns(4)
                agenda_nova[d] = {
                    "HTPC": [cols[0].text_input("HTPC Ini", key=f"hi_{d}"), cols[0].text_input("HTPC Fim", key=f"hf_{d}")],
                    "HSP1": [cols[1].text_input("HSP Ini", key=f"h1i_{d}"), cols[1].text_input("HSP Fim", key=f"h1f_{d}")],
                    "HSP2": [cols[2].text_input("HSP Ini", key=f"h2i_{d}"), cols[2].text_input("HSP Fim", key=f"h2f_{d}")],
                    "HE": [cols[3].text_input("HE Ini", key=f"hei_{d}"), cols[3].text_input("HE Fim", key=f"hef_{d}")]
                }
        if st.form_submit_button("SALVAR CADASTRO"):
            if nome:
                st.session_state.profs[nome] = {"nome":nome, "rg":rg, "situacao":situacao, "disciplina":disciplina, "categoria":categoria, "agenda":agenda_nova, "ativo":ativo}
                save_data(DB_FILE, st.session_state.profs); st.success(f"Cadastrado!")
            else: st.error("Insira o nome.")

# 2. ABA EDITAR
with tabs[1]:
    st.subheader("Manutenção de Dados")
    prof_nome = st.selectbox("Selecione o Professor", [""] + list(st.session_state.profs.keys()))
    if prof_nome:
        p = st.session_state.profs[prof_nome]
        with st.form("edit_form"):
            e1, e2, e3 = st.columns(3); enome = e1.text_input("Nome", value=p['nome']); erg = e2.text_input("RG", value=p['rg']); esit = e3.selectbox("Situação", ["ACTS", "EFETIVO"], index=0 if p['situacao']=="ACTS" else 1)
            e4, e5, e6 = st.columns([2,2,1]); edisc = e4.text_input("Disciplina", value=p['disciplina']); ecat = e5.text_input("Categoria", value=p['categoria']); eativo = e6.checkbox("Professor Ativo", value=p.get('ativo', True))
            e_agenda = {}
            for d in ["Segunda", "Terça", "Quarta", "Quinta", "Sexta"]:
                with st.expander(f"Editar {d}"):
                    cols = st.columns(4); h = p['agenda'].get(d, {"HTPC":["",""],"HSP1":["",""],"HSP2":["",""],"HE":["",""]})
                    e_agenda[d] = {
                        "HTPC": [cols[0].text_input("HTPC Ini", value=h['HTPC'][0], key=f"ehi_{d}"), cols[0].text_input("HTPC Fim", value=h['HTPC'][1], key=f"ehf_{d}")],
                        "HSP1": [cols[1].text_input("HSP Ini", value=h['HSP1'][0], key=f"eh1i_{d}"), cols[1].text_input("HSP Fim", value=h['HSP1'][1], key=f"eh1f_{d}")],
                        "HSP2": [cols[2].text_input("HSP Ini", value=h['HSP2'][0], key=f"eh2i_{d}"), cols[2].text_input("HSP Fim", value=h['HSP2'][1], key=f"eh2f_{d}")],
                        "HE": [cols[3].text_input("HE Ini", value=h['HE'][0], key=f"ehei_{d}"), cols[3].text_input("HE Fim", value=h['HE'][1], key=f"hef_{d}")]
                    }
            if st.form_submit_button("ATUALIZAR DADOS"):
                if enome != prof_nome: del st.session_state.profs[prof_nome]
                st.session_state.profs[enome] = {"nome":enome, "rg":erg, "situacao":esit, "disciplina":edisc, "categoria":ecat, "agenda":e_agenda, "ativo":eativo}
                save_data(DB_FILE, st.session_state.profs); st.success("Atualizado!"); st.rerun()
        st.markdown('<div class="btn-excluir">', unsafe_allow_html=True)
        if st.button("EXCLUIR PROFESSOR"):
            del st.session_state.profs[prof_nome]; save_data(DB_FILE, st.session_state.profs); st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

# 3. ABA GERAR RELATÓRIO
with tabs[2]:
    st.subheader("Geração de Lote Mensal")
    nlote = st.text_input("Identificação do Lote"); r1, r2 = st.columns(2); d1 = r1.date_input("Início do Período", value=datetime(2026, 3, 15), format="DD/MM/YYYY"); d2 = r2.date_input("Fim do Período", value=datetime(2026, 4, 13), format="DD/MM/YYYY")
    st.write("---"); profs_ativos = {k: v for k, v in st.session_state.profs.items() if v.get('ativo', True)}
    def toggle_all():
        for pn in profs_ativos.keys(): st.session_state[f"sel_{pn}"] = st.session_state.all_active
    st.checkbox("Selecionar todos os professores ativos", key="all_active", on_change=toggle_all)
    p_nomes_lote = []
    for pn in profs_ativos.keys():
        if st.checkbox(pn, key=f"sel_{pn}"): p_nomes_lote.append(pn)
    if st.button("GERAR RELATÓRIOS"):
        if nlote and p_nomes_lote:
            lid = f"{nlote}_{datetime.now().strftime('%Y%m%d%H%M%S')}"; st.session_state.lotes[lid] = {"nome":nlote, "inicio":str(d1), "fim":str(d2), "inicio_f":d1.strftime("%d/%m/%Y"), "fim_f":d2.strftime("%d/%m/%Y"), "professores":{n: profs_ativos[n] for n in p_nomes_lote}}
            save_data(HIST_FILE, st.session_state.lotes); st.success("Lote pronto no Histórico!")

# 4. ABA HISTÓRICO
with tabs[3]:
    st.subheader("Documentos Arquivados")
    if not st.session_state.lotes: st.info("Vazio.")
    else:
        for lid, ld in reversed(list(st.session_state.lotes.items())):
            with st.expander(f"📁 {ld['nome']} | {ld['inicio_f']} a {ld['fim_f']}"):
                col_h1, col_h2 = st.columns([1, 4]); w_buf = gerar_docx_modelo(ld)
                col_h1.download_button("📥 Baixar Word", data=w_buf, file_name=f"{ld['nome']}.docx", key=f"dl_{lid}")
                if col_h2.button("🗑️ Excluir Lote", key=f"del_h_{lid}"):
                    del st.session_state.lotes[lid]; save_data(HIST_FILE, st.session_state.lotes); st.rerun()